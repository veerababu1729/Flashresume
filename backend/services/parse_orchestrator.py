import os
import io
import gc
import pypdfium2 as pdfium  # Replaced PyMuPDF with pypdfium2 for better stability
from docx import Document
from services.pdf_parser import extract_with_pdfplumber, is_extraction_good
from services.pdf_link_extractor import extract_pdf_links
from services.docx_link_extractor import extract_docx_links


def extract_from_docx(docx_bytes: bytes) -> dict:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(paragraphs)
        
        extracted_links = extract_docx_links(doc, resume_text=text)
        
        return {
            "text": text if text else "Unable to extract text from DOCX.",
            "page_count": 1,
            "parser_used": "python-docx",
            "extracted_links": extracted_links,
        }
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_with_pypdfium2(pdf_bytes: bytes) -> tuple[str, int]:
    """Layer 1: pypdfium2 — Fast C++ engine, lightweight memory (~15MB)."""
    doc = pdfium.PdfDocument(pdf_bytes)
    page_count = len(doc)
    text_parts = []
    for page in doc:
        text_page = page.get_textpage()
        text_parts.append(text_page.get_text_bounded().strip())
    return "\n\n".join(text_parts), page_count


def extract_resume_text(pdf_bytes: bytes) -> dict:
    """
    2-layer orchestrator:
    Layer 1: pypdfium2 (Fast, lightweight C++ engine)
    Layer 2: pdfplumber (Heavy Python heap, fallback)
    Note: OCR fallback was removed to prevent OOM memory issues on the server.

    Additionally runs pdf_link_extractor in-process to harvest real hyperlink
    URLs from PDF annotation objects (invisible to plain-text extraction).
    """
    # Layer 1 — pypdfium2
    text, page_count = extract_with_pypdfium2(pdf_bytes)
    if is_extraction_good(text):
        # Run link extractor on the same PDF bytes (annotation walk + regex on text)
        extracted_links = extract_pdf_links(pdf_bytes, resume_text=text)
        return {
            "text": text,
            "page_count": page_count,
            "parser_used": "pypdfium2",
            "extracted_links": extracted_links,
        }

    # Layer 2 — pdfplumber (Fallback)
    try:
        text, page_count = extract_with_pdfplumber(pdf_bytes)
        if is_extraction_good(text):
            # Still run the link extractor — annotation walk doesn't depend on text quality
            extracted_links = extract_pdf_links(pdf_bytes, resume_text=text)
            return {
                "text": text,
                "page_count": page_count,
                "parser_used": "pdfplumber",
                "extracted_links": extracted_links,
            }
    finally:
        gc.collect()

    raise ValueError("We couldn't read the text in this file. Please upload a standard digital PDF or Word document instead of a scanned image.")
